import os
import sys
import subprocess

# Auto-install python-docx if not present
try:
    import docx
except ImportError:
    print("Installing python-docx library...")
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
    import docx

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, fill_color):
    """Sets background color for table cells."""
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_color}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    """Sets cell padding."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('w:top', top), ('w:bottom', bottom), ('w:left', left), ('w:right', right)]:
        node = OxmlElement(m)
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def add_border_to_table(table, color="CCCCCC"):
    """Adds horizontal and vertical borders to a table."""
    tblPr = table._tbl.tblPr
    tblBorders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="4" w:space="0" w:color="{color}"/>'
        f'  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="{color}"/>'
        f'  <w:left w:val="none"/>'
        f'  <w:right w:val="none"/>'
        f'  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="{color}"/>'
        f'  <w:insideV w:val="single" w:sz="4" w:space="0" w:color="{color}"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(tblBorders)

def add_code_block_shading(paragraph, color="F4F4F4"):
    """Adds light grey background fill to code blocks."""
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:left w:val="single" w:sz="24" w:space="8" w:color="002C9B"/>'
        f'</w:pBdr>'
    )
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    pPr.append(pBdr)
    pPr.append(shd)

def build_document():
    md_path = "Strivo_Project_Documentation.md"
    docx_path = "Strivo_Project_Documentation.docx"
    
    if not os.path.exists(md_path):
        print(f"Error: {md_path} not found.")
        return

    doc = Document()
    
    # Configure 1 inch page margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Configure styles
    style_normal = doc.styles['Normal']
    font = style_normal.font
    font.name = 'Arial'
    font.size = Pt(10.5)
    style_normal.paragraph_format.line_spacing = 1.15
    style_normal.paragraph_format.space_after = Pt(6)

    # Palette
    NAVY = RGBColor(0, 44, 155)      # Primary Accent: #002C9B
    CHARCOAL = RGBColor(51, 51, 51)  # Body text
    GREY = RGBColor(128, 128, 128)

    with open(md_path, "r", encoding="utf-8") as f:
        lines = f.readlines()

    in_code_block = False
    code_block_lines = []
    
    in_table = False
    table_rows = []

    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # Handle code blocks
        if stripped.startswith("```"):
            if in_code_block:
                # Close code block
                code_text = "\n".join(code_block_lines)
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.2)
                p.paragraph_format.space_before = Pt(4)
                p.paragraph_format.space_after = Pt(6)
                add_code_block_shading(p)
                
                run = p.add_run(code_text)
                run.font.name = 'Courier New'
                run.font.size = Pt(8.5)
                run.font.color.rgb = CHARCOAL
                
                code_block_lines = []
                in_code_block = False
            else:
                in_code_block = True
            i += 1
            continue

        if in_code_block:
            code_block_lines.append(line.rstrip('\n'))
            i += 1
            continue

        # Handle tables
        if stripped.startswith("|"):
            # Check if it is a separator line (e.g. |:---| or |--|)
            if i + 1 < len(lines) and all(c in '|:- ' for c in lines[i+1].strip()):
                in_table = True
                # Parse headers
                headers = [cell.strip() for cell in stripped.split("|")[1:-1]]
                table_rows.append(headers)
                i += 2  # skip separator line
                continue
            elif in_table:
                # Add data row
                cells = [cell.strip() for cell in stripped.split("|")[1:-1]]
                table_rows.append(cells)
                i += 1
                continue
        
        # If table was building and now ended
        if in_table and not stripped.startswith("|"):
            # Create table
            num_cols = len(table_rows[0])
            table = doc.add_table(rows=len(table_rows), cols=num_cols)
            table.autofit = True
            add_border_to_table(table, "D3D3D3")
            
            for row_idx, row_data in enumerate(table_rows):
                row = table.rows[row_idx]
                is_header = (row_idx == 0)
                
                for col_idx, cell_value in enumerate(row_data):
                    if col_idx >= len(row.cells):
                        break
                    cell = row.cells[col_idx]
                    cell.text = ""  # Clear default text
                    
                    p = cell.paragraphs[0]
                    p.paragraph_format.space_after = Pt(2)
                    p.paragraph_format.space_before = Pt(2)
                    
                    # Highlight bold texts in cells
                    run = p.add_run(cell_value.replace("**", ""))
                    run.font.name = 'Arial'
                    run.font.size = Pt(9)
                    
                    if is_header:
                        run.bold = True
                        run.font.color.rgb = RGBColor(255, 255, 255)
                        set_cell_background(cell, "002C9B") # Navy header
                    else:
                        run.font.color.rgb = CHARCOAL
                        if "**" in cell_value:
                            run.bold = True
                        if row_idx % 2 == 0:
                            set_cell_background(cell, "F9FBFD") # Striping
                        else:
                            set_cell_background(cell, "FFFFFF")
                            
                    set_cell_margins(cell, top=80, bottom=80, left=120, right=120)
            
            # Spacer after table
            doc.add_paragraph()
            table_rows = []
            in_table = False
            # Note: do not skip this line, it will be evaluated in the next loop iteration
            continue

        # Skip separator lines
        if stripped == "---":
            # Add a visual horizontal rule spacer
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(12)
            p_xml = p._p.get_or_add_pPr()
            pbdr = parse_xml(f'<w:pBdr {nsdecls("w")}><w:bottom w:val="single" w:sz="6" w:space="1" w:color="002C9B"/></w:pBdr>')
            p_xml.append(pbdr)
            i += 1
            continue

        # Handle headings
        if stripped.startswith("<h1") or stripped.startswith("<h2"):
            # HTML centered headings
            # Extract text content between HTML tags
            import re
            text_match = re.search(r'>(.*?)<', stripped)
            text = text_match.group(1).strip() if text_match else stripped
            
            # Determine Heading level
            is_title = "h1" in stripped
            
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(18)
            p.paragraph_format.space_after = Pt(8)
            p.paragraph_format.keep_with_next = True
            
            run = p.add_run(text)
            run.bold = True
            run.font.color.rgb = NAVY
            run.font.size = Pt(20 if is_title else 14)
            
            i += 1
            continue

        if stripped.startswith("###"):
            text = stripped.replace("###", "").strip()
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.keep_with_next = True
            
            run = p.add_run(text)
            run.bold = True
            run.font.color.rgb = NAVY
            run.font.size = Pt(11)
            
            i += 1
            continue

        if stripped.startswith("####"):
            text = stripped.replace("####", "").strip()
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.keep_with_next = True
            
            run = p.add_run(text)
            run.bold = True
            run.font.color.rgb = CHARCOAL
            run.font.size = Pt(10)
            
            i += 1
            continue

        # Handle bullet points
        if stripped.startswith("- ") or stripped.startswith("* ") or stripped.startswith("• "):
            text = stripped[2:].strip()
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.space_after = Pt(3)
            
            # Formatting inline bold variables (e.g. - **Title**: description)
            if "**" in text:
                parts = text.split("**")
                for idx, part in enumerate(parts):
                    run = p.add_run(part)
                    run.font.color.rgb = CHARCOAL
                    if idx % 2 == 1:
                        run.bold = True
            else:
                run = p.add_run(text)
                run.font.color.rgb = CHARCOAL
            
            i += 1
            continue

        # Paragraph text
        if stripped:
            p = doc.add_paragraph()
            
            # Format bold text
            if "**" in stripped:
                parts = stripped.split("**")
                for idx, part in enumerate(parts):
                    run = p.add_run(part)
                    run.font.color.rgb = CHARCOAL
                    if idx % 2 == 1:
                        run.bold = True
            else:
                run = p.add_run(stripped)
                run.font.color.rgb = CHARCOAL

        i += 1

    doc.save(docx_path)
    print(f"Successfully generated {docx_path}!")

if __name__ == "__main__":
    build_document()
